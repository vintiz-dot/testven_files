import docx
import sys

def read_docx(path):
    try:
        doc = docx.Document(path)
        with open("docx_out.txt", "w", encoding="utf-8") as f:
            for para in doc.paragraphs:
                f.write(para.text + "\n")
            
            for table in doc.tables:
                for row in table.rows:
                    row_data = []
                    for cell in row.cells:
                        row_data.append(cell.text.replace("\n", " "))
                    f.write(" | ".join(row_data) + "\n")
    except Exception as e:
        print("Error:", e)

if __name__ == '__main__':
    read_docx(sys.argv[1])
